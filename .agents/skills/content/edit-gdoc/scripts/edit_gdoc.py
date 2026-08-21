#!/usr/bin/env python3
"""Edit a Google Doc by exporting as DOCX, modifying paragraphs by index, and re-uploading.

Preserves images, tables, and formatting. Default scope is drive.file (files this
OAuth client created). Auth is Application Default Credentials.
"""

import argparse
import json
import os
import re
import sys
import tempfile

import google.auth
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload

from docx import Document

SCOPES = ["https://www.googleapis.com/auth/drive.file"]
TEMP_DIR = os.path.join(tempfile.gettempdir(), "edit-gdoc")


def authenticate():
    creds, _ = google.auth.default(scopes=SCOPES)
    if not creds.valid:
        creds.refresh(Request())
    return creds


def extract_doc_id(url_or_id: str) -> str:
    match = re.search(r"/document/d/([a-zA-Z0-9_-]+)", url_or_id)
    return match.group(1) if match else url_or_id


def download_docx(drive_service, doc_id: str) -> str:
    """Always pull the latest version. Returns path to fresh .docx."""
    os.makedirs(TEMP_DIR, exist_ok=True)
    path = os.path.join(TEMP_DIR, f"{doc_id}.docx")

    if os.path.exists(path):
        os.remove(path)

    request = drive_service.files().export_media(
        fileId=doc_id,
        mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    with open(path, "wb") as f:
        downloader = MediaIoBaseDownload(f, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()

    return path


def upload_docx(drive_service, doc_id: str, docx_path: str):
    media = MediaFileUpload(
        docx_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    drive_service.files().update(fileId=doc_id, media_body=media).execute()


def cmd_pull(args):
    """Pull latest and display paragraphs with indices."""
    creds = authenticate()
    drive_service = build("drive", "v3", credentials=creds)
    doc_id = extract_doc_id(args.doc)

    docx_path = download_docx(drive_service, doc_id)
    doc = Document(docx_path)

    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "style": para.style.name if para.style else "Normal",
                "text": text,
            })

    if args.json:
        print(json.dumps(paragraphs, indent=2))
    elif args.full:
        for p in paragraphs:
            print(f"\n--- Paragraph {p['index']} [{p['style']}] ---")
            print(p["text"])
        print(f"\n{len(paragraphs)} paragraphs total.")
    else:
        print(f"Document has {len(paragraphs)} text paragraphs:\n")
        for p in paragraphs:
            preview = p["text"][:120]
            if len(p["text"]) > 120:
                preview += "..."
            print(f"  [{p['index']:>3}] ({p['style']}) {preview}")

    if os.path.exists(docx_path):
        os.remove(docx_path)


def cmd_replace(args):
    """Pull latest, replace paragraph at index, upload."""
    creds = authenticate()
    drive_service = build("drive", "v3", credentials=creds)
    doc_id = extract_doc_id(args.doc)

    docx_path = download_docx(drive_service, doc_id)
    doc = Document(docx_path)

    if args.paragraph < 0 or args.paragraph >= len(doc.paragraphs):
        print(
            f"Error: index {args.paragraph} out of range (0-{len(doc.paragraphs)-1}).",
            file=sys.stderr,
        )
        os.remove(docx_path)
        sys.exit(1)

    target = doc.paragraphs[args.paragraph]
    old_text = target.text

    print(f"Replacing paragraph {args.paragraph}:")
    print(f"  OLD: {old_text[:120]}{'...' if len(old_text) > 120 else ''}")
    print(f"  NEW: {args.text[:120]}{'...' if len(args.text) > 120 else ''}")

    for run in target.runs[1:]:
        run.text = ""
    if target.runs:
        target.runs[0].text = args.text
    else:
        target.text = args.text

    doc.save(docx_path)
    upload_docx(drive_service, doc_id, docx_path)
    os.remove(docx_path)
    print("Done.")


def main():
    parser = argparse.ArgumentParser(
        description="Edit a Google Doc by paragraph index via DOCX roundtrip"
    )
    sub = parser.add_subparsers(dest="command", required=True)

    pull = sub.add_parser("pull", help="Pull latest and display paragraphs")
    pull.add_argument("doc", help="Google Doc URL or ID")
    pull.add_argument("--json", action="store_true", help="Output as JSON")
    pull.add_argument("--full", action="store_true", help="Show full paragraph text")

    replace = sub.add_parser("replace", help="Replace a paragraph by index")
    replace.add_argument("doc", help="Google Doc URL or ID")
    replace.add_argument("--paragraph", "-p", type=int, required=True, help="Paragraph index")
    replace.add_argument("--text", "-t", required=True, help="New paragraph text")

    args = parser.parse_args()
    if args.command == "pull":
        cmd_pull(args)
    elif args.command == "replace":
        cmd_replace(args)


if __name__ == "__main__":
    main()
